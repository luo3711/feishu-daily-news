#!/usr/bin/env python3
"""Fuel Cell Intelligence Daily - DOCX Edition"""

import json, os, urllib.request, urllib.parse, base64, hmac, hashlib
import xml.etree.ElementTree as ET
import time as time_module
from datetime import datetime, timedelta
from io import BytesIO

WEBHOOK_URL   = os.environ["FEISHU_WEBHOOK_URL"]
FEISHU_SECRET = os.environ["FEISHU_SECRET"]
GITHUB_TOKEN  = os.environ["GITHUB_TOKEN"]
REPO          = os.environ.get("GITHUB_REPOSITORY", "bake7791/feishu-daily-news")
AI_ENDPOINT   = "https://models.inference.ai.azure.com/chat/completions"
AI_MODEL      = "gpt-4o-mini"

QUERIES = [
    ("fuel cell vehicle policy hydrogen regulation 2026", "en-US", "US"),
    ("hydrogen fuel cell FCEV industry government strategy", "en-GB", "GB"),
    ("Brennstoffzelle Wasserstoff Fahrzeug EU Politik", "de-DE", "DE"),
    ("pile combustible hydrogene vehicule politique France", "fr-FR", "FR"),
]
CN_QUERIES = [
    ("fuel cell vehicle China policy hydrogen", "zh-CN", "CN"),
    ("hydrogen fuel cell industry news China", "zh-CN", "CN"),
]
JP_QUERIES = [
    ("fuel cell vehicle Japan Toyota Honda policy", "ja-JP", "JP"),
]

MAX_ARTICLES = 30


def http_get(url, headers=None):
    req = urllib.request.Request(url, headers=headers or {})
    with urllib.request.urlopen(req, timeout=30) as r:
        return r.read().decode("utf-8")


def http_post_json(url, payload, headers=None):
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    h = {"Content-Type": "application/json; charset=utf-8"}
    if headers: h.update(headers)
    req = urllib.request.Request(url, data=data, headers=h, method="POST")
    with urllib.request.urlopen(req, timeout=90) as r:
        return json.loads(r.read().decode("utf-8"))


def search_google_news(query, hl, gl, max_results=50):
    ceid_map = {"CN":"CN:zh-Hans","US":"US:en","GB":"GB:en","JP":"JP:ja","DE":"DE:de","FR":"FR:fr"}
    ceid = ceid_map.get(gl, f'{gl}:{hl.split("-")[0]}')
    rss_url = f"https://news.google.com/rss/search?q={urllib.parse.quote(query)}&hl={hl}&gl={gl}&ceid={ceid}"
    xml_data = http_get(rss_url, {"User-Agent": "Mozilla/5.0"})
    root = ET.fromstring(xml_data)
    results = []
    for item in root.findall(".//item"):
        t = item.find("title"); l = item.find("link"); s = item.find("source"); p = item.find("pubDate")
        title = t.text.strip() if t is not None and t.text else ""
        link = l.text if l is not None else ""
        source = s.text.strip() if s is not None and s.text else "Unknown"
        pubdate = p.text if p is not None else ""
        skip = ["stock", "share price", "sponsored"]
        if not title or any(w in title.lower() for w in skip):
            continue
        results.append({"title":title,"url":link,"source":source,"date":pubdate,"region":gl})
        if len(results)>=max_results: break
    return results


def search_all():
    seen = set(); all_results = []
    for query, hl, gl in QUERIES + CN_QUERIES + JP_QUERIES:
        try:
            results = search_google_news(query, hl, gl)
            print(f"  [{gl}] -> {len(results)} results")
            for r in results:
                key = r["title"][:80]
                if key not in seen:
                    seen.add(key); all_results.append(r)
        except Exception as e:
            print(f"  [WARN] {gl}: {e}")
    def pd(r):
        try: return datetime.strptime(r["date"], "%a, %d %b %Y %H:%M:%S %Z")
        except: return datetime.min
    all_results.sort(key=pd, reverse=True)
    return all_results[:MAX_ARTICLES]


def ai_analyze(articles):
    articles_text = ""
    for i, a in enumerate(articles):
        flag = {"CN":"[CN]","US":"[US]","GB":"[UK]","JP":"[JP]","DE":"[DE]","FR":"[FR]"}.get(a["region"],"")
        articles_text += f"\n{i+1}. {flag} {a['title']} (source: {a['source']})"

    prompt = f"""You are a fuel cell vehicle industry analyst. Below are today's latest global news ({len(articles)} items). Write a comprehensive analysis report in Chinese, using Markdown format:

{articles_text}

Structure:
## 1. Today's Key Headlines (pick 3-5 most important, ~60 chars each in Chinese summary)
## 2. Policy & Regulation Dynamics (latest policy changes, subsidies, regulations across countries)
## 3. Industry & Technology Trends (tech breakthroughs, production progress, supply chain)
## 4. Strategic Assessment (~100 chars brief judgment on industry trend)

Be professional and concise."""

    result = http_post_json(AI_ENDPOINT, {
        "model": AI_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2000,
        "temperature": 0.3,
    }, headers={"Authorization": f"Bearer {GITHUB_TOKEN}"})
    return result["choices"][0]["message"]["content"]


def build_docx(ai_report, articles):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading("Fuel Cell Vehicle Intelligence Daily", level=0)
    sub = doc.add_paragraph(f"Date: {today}  |  Global Sources  |  GPT-4o-mini AI Analysis")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50)

    sections = ai_report.split("## ")
    for sec in sections:
        sec = sec.strip()
        if not sec: continue
        lines = sec.split("\n", 1)
        h_text = lines[0].strip()
        body = lines[1].strip() if len(lines) > 1 else ""
        doc.add_heading(h_text, level=2)
        if body:
            doc.add_paragraph(body)

    doc.add_paragraph("-" * 50)
    doc.add_heading("Sources", level=2)

    for i, a in enumerate(articles, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {a["title"]}')
        run.bold = True; run.font.size = Pt(10)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f'    Link: {a["url"]}')
        r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0, 102, 204)
        p3 = doc.add_paragraph(f'    Source: {a["source"]}  |  {a["date"][:22]}')
        p3.runs[0].font.size = Pt(8)

    doc.add_paragraph("")
    doc.add_paragraph("-" * 50)
    doc.add_paragraph("Auto-generated by GitHub Actions. For reference only.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def upload_release(docx_buf, date_str):
    filename = f"fuel_cell_report_{date_str}.docx"
    api_base = f"https://api.github.com/repos/{REPO}"
    headers = {"Authorization": f"Bearer {GITHUB_TOKEN}", "Accept": "application/vnd.github+json"}

    try:
        release = json.loads(http_get(f"{api_base}/releases/tags/daily-report", headers=headers))
        release_id = release["id"]
    except:
        release = http_post_json(f"{api_base}/releases", {
            "tag_name": "daily-report", "name": "Daily Reports",
            "body": "Fuel Cell Intelligence Daily DOCX Reports"
        }, headers=headers)
        release_id = release["id"]

    upload_url = f"https://uploads.github.com/repos/{REPO}/releases/{release_id}/assets?name={urllib.parse.quote(filename)}"
    uh = {"Authorization": f"Bearer {GITHUB_TOKEN}", "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Accept": "application/vnd.github+json"}
    req = urllib.request.Request(upload_url, data=docx_buf.read(), headers=uh, method="POST")
    with urllib.request.urlopen(req, timeout=30) as r:
        asset = json.loads(r.read().decode())
    print(f"  Uploaded: {asset['browser_download_url']}")
    return asset["browser_download_url"]


def send_feishu(download_url, ai_summary, today):
    summary = ai_summary[:250].replace("#","").replace("*","").strip() + "..."

    card_text = f"""**Fuel Cell Intelligence Daily** - {today}

{summary}

---

**Download Full Report (DOCX):** [{download_url}]({download_url})

> GPT-4o-mini AI Analysis | Google News Global Search
> Daily 08:00 (Beijing Time) Auto Push"""

    payload = {
        "msg_type": "interactive",
        "card": {
            "header": {"title": {"tag": "plain_text", "content": f"Fuel Cell Daily - {today}"}, "template": "blue"},
            "elements": [{"tag": "markdown", "content": card_text}],
        },
    }

    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    for attempt in range(3):
        ts = str(int(time_module.time()))
        sk = (ts + "\n" + FEISHU_SECRET).encode("utf-8")
        sig = base64.b64encode(hmac.new(sk, b"", hashlib.sha256).digest()).decode()
        url = f"{WEBHOOK_URL}?timestamp={ts}&sign={sig}"
        r = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json; charset=utf-8"}, method="POST")
        try:
            with urllib.request.urlopen(r, timeout=15) as resp:
                result = json.loads(resp.read().decode())
            if result.get("code") == 0:
                print(f"[OK] Pushed (attempt {attempt+1})")
                return
            print(f"[RETRY {attempt+1}] {result}")
            time_module.sleep(2)
        except Exception as e:
            print(f"[RETRY {attempt+1}] {e}")
            time_module.sleep(2)
    raise Exception("Push failed")


def main():
    print("=" * 50)
    print("Fuel Cell Intelligence Daily - DOCX")
    print("=" * 50)
    td = datetime.now()
    today = td.strftime("%Y%m%d")
    today_cn = td.strftime("%m.%d")

    print("\n[1/4] Searching...")
    articles = search_all()
    print(f"  Total: {len(articles)}")
    if not articles: return send_feishu("", "No news", today_cn)

    print("\n[2/4] AI Analyzing...")
    try:
        ai_report = ai_analyze(articles)
        print(f"  {len(ai_report)} chars")
    except Exception as e:
        ai_report = f"Found {len(articles)} articles. See DOCX for details."
        print(f"  AI failed: {e}")

    print("\n[3/4] Building DOCX...")
    docx_buf = build_docx(ai_report, articles)

    print("\n[4/4] Uploading & Pushing...")
    url = upload_release(docx_buf, today)
    send_feishu(url, ai_report, today_cn)
    print(f"\nDone! {url}")


if __name__ == "__main__":
    main()